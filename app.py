import streamlit as st
import pandas as pd
import zipfile
import rarfile
import io
import re
from collections import defaultdict
import matplotlib.pyplot as plt
import seaborn as sns
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image, PageBreak
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, landscape
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
import tempfile
import os
from datetime import datetime
from docx import Document
from docx.shared import Inches
import json
from datetime import timedelta

st.set_page_config(
    page_title="Network Config Audit",
    page_icon="🛡️",
    layout="centered"
)

# ---------------------------
# Audit Planning Assistant
# ---------------------------
def audit_planner():
    st.header("📅 Network Audit Planning Assistant")
    
    # Audit Scope
    st.subheader("1. Audit Scope Definition")
    col1, col2 = st.columns(2)
    
    with col1:
        audit_name = st.text_input("Audit Name", "Q1 2024 Network Security Audit")
        audit_type = st.selectbox(
            "Audit Type",
            ["Comprehensive Security", "Compliance Check", "Pre-Migration", "Post-Change", "Routine Maintenance"]
        )
    
    with col2:
        priority = st.select_slider("Priority Level", ["Low", "Medium", "High", "Critical"])
        timeline_days = st.number_input("Timeline (days)", min_value=1, max_value=90, value=14)
    
    # Device Inventory
    st.subheader("2. Device Inventory")
    
    device_types = st.multiselect(
        "Device Types to Audit",
        ["Routers", "Switches", "Firewalls", "Wireless Controllers", "Load Balancers", "VPN Gateways"],
        default=["Routers", "Switches", "Firewalls"]
    )
    
    estimated_devices = st.number_input("Estimated Number of Devices", min_value=1, max_value=1000, value=50)
    
    # Risk Assessment
    st.subheader("3. Risk Assessment Factors")
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        business_impact = st.selectbox(
            "Business Impact",
            ["Low", "Medium", "High", "Critical"],
            help="Impact on business operations if devices fail"
        )
    
    with col2:
        data_sensitivity = st.selectbox(
            "Data Sensitivity",
            ["Public", "Internal", "Confidential", "Restricted"],
            help="Sensitivity of data handled by these devices"
        )
    
    with col3:
        compliance_requirements = st.multiselect(
            "Compliance Requirements",
            ["PCI-DSS", "HIPAA", "SOX", "GDPR", "NIST", "ISO 27001", "None"]
        )
    
    # Resource Planning
    st.subheader("4. Resource Planning")
    
    col1, col2 = st.columns(2)
    
    with col1:
        team_size = st.number_input("Team Size", min_value=1, max_value=20, value=3)
        hours_per_device = st.slider("Estimated Hours per Device", 0.5, 8.0, 2.0)
    
    with col2:
        expertise_level = st.selectbox(
            "Required Expertise Level",
            ["Junior", "Mid-Level", "Senior", "Expert"]
        )
        tools_available = st.multiselect(
            "Available Tools",
            ["Network Scanner", "Config Manager", "SIEM", "Vulnerability Scanner", "Custom Scripts"]
        )
    
    # Timeline Planning
    st.subheader("5. Timeline & Milestones")
    
    start_date = st.date_input("Planned Start Date", datetime.now() + timedelta(days=7))
    
    # Calculate timeline
    total_hours = estimated_devices * hours_per_device
    total_days = max(1, total_hours / (team_size * 8))  # 8 hours per day per person
    
    # Key milestones - convert dates to strings for JSON serialization
    milestones = {
        "Planning & Scoping": start_date.strftime("%Y-%m-%d"),
        "Data Collection": (start_date + timedelta(days=2)).strftime("%Y-%m-%d"),
        "Configuration Analysis": (start_date + timedelta(days=int(total_days * 0.3))).strftime("%Y-%m-%d"),
        "Vulnerability Assessment": (start_date + timedelta(days=int(total_days * 0.6))).strftime("%Y-%m-%d"),
        "Reporting": (start_date + timedelta(days=int(total_days * 0.8))).strftime("%Y-%m-%d"),
        "Remediation Planning": (start_date + timedelta(days=int(total_days))).strftime("%Y-%m-%d")
    }
    
    # Display planning summary
    if st.button("Generate Audit Plan"):
        st.success("🎯 Audit Plan Generated Successfully!")
        
        # Summary Section
        st.subheader("📋 Audit Plan Summary")
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            st.metric("Total Devices", estimated_devices)
            st.metric("Team Size", team_size)
            st.metric("Priority", priority)
        
        with col2:
            st.metric("Total Hours", f"{total_hours:.1f}")
            st.metric("Estimated Days", f"{total_days:.1f}")
            st.metric("Business Impact", business_impact)
        
        with col3:
            st.metric("Start Date", start_date.strftime("%Y-%m-%d"))
            st.metric("End Date", (start_date + timedelta(days=total_days)).strftime("%Y-%m-%d"))
            st.metric("Risk Level", "High" if business_impact in ["High", "Critical"] else "Medium")
        
        # Timeline Visualization
        st.subheader("⏰ Project Timeline")
        timeline_data = []
        for milestone, date_str in milestones.items():
            date = datetime.strptime(date_str, "%Y-%m-%d").date()
            timeline_data.append({
                "Milestone": milestone,
                "Date": date_str,
                "Days from Start": (date - start_date).days
            })
        
        timeline_df = pd.DataFrame(timeline_data)
        st.dataframe(timeline_df, width='stretch')
        
        # Resource Allocation
        st.subheader("👥 Resource Allocation")
        
        resource_data = {
            "Task": ["Planning", "Data Collection", "Analysis", "Reporting", "Remediation"],
            "Effort (%)": [10, 25, 40, 15, 10],
            "Team Members": [team_size, team_size, team_size, team_size - 1, team_size]
        }
        resource_df = pd.DataFrame(resource_data)
        st.dataframe(resource_df, width='stretch')
        
        # Risk Matrix
        st.subheader("🚨 Risk Assessment Matrix")
        
        risk_matrix = {
            "Factor": ["Business Impact", "Data Sensitivity", "Compliance", "Team Expertise", "Tool Availability"],
            "Level": [business_impact, data_sensitivity, 
                     "High" if compliance_requirements else "Low", 
                     expertise_level,
                     "High" if len(tools_available) >= 3 else "Medium"],
            "Mitigation": [
                "Ensure backup systems available",
                "Focus on encryption & access controls",
                "Document compliance evidence",
                "Provide training if needed",
                "Plan for manual processes"
            ]
        }
        risk_df = pd.DataFrame(risk_matrix)
        st.dataframe(risk_df, width='stretch')
        
        # Export Plan
        st.subheader("📤 Export Audit Plan")
        
        audit_plan = {
            "audit_name": audit_name,
            "audit_type": audit_type,
            "priority": priority,
            "timeline_days": timeline_days,
            "device_types": device_types,
            "estimated_devices": estimated_devices,
            "team_size": team_size,
            "total_hours": total_hours,
            "start_date": start_date.strftime("%Y-%m-%d"),
            "milestones": milestones,
            "risk_factors": risk_matrix
        }
        
        # JSON export
        json_plan = json.dumps(audit_plan, indent=2)
        st.download_button(
            label="Download Audit Plan (JSON)",
            data=json_plan,
            file_name=f"audit_plan_{datetime.now().strftime('%Y%m%d')}.json",
            mime="application/json"
        )
        
        # Text summary export
        text_summary = f"""
AUDIT PLAN: {audit_name}
==================================
Type: {audit_type}
Priority: {priority}
Timeline: {timeline_days} days
Start Date: {start_date.strftime('%Y-%m-%d')}

DEVICE INVENTORY:
-----------------
Types: {', '.join(device_types)}
Estimated Devices: {estimated_devices}

RESOURCE PLANNING:
------------------
Team Size: {team_size}
Expertise Level: {expertise_level}
Total Effort: {total_hours} hours
Estimated Duration: {total_days:.1f} days

RISK ASSESSMENT:
----------------
Business Impact: {business_impact}
Data Sensitivity: {data_sensitivity}
Compliance: {', '.join(compliance_requirements) if compliance_requirements else 'None'}

MILESTONES:
-----------
{chr(10).join([f'{milestone}: {date_str}' for milestone, date_str in milestones.items()])}
        """
        
        st.download_button(
            label="Download Audit Summary (TXT)",
            data=text_summary,
            file_name=f"audit_summary_{datetime.now().strftime('%Y%m%d')}.txt",
            mime="text/plain"
        )

# ---------------------------
# Audit function (7 categories)
# returns list of tuples (Finding, File, RiskDesc, Recommendation, Category)
# ---------------------------
def audit_config(filename, content):
    findings = []

    # Normalize content for easier regex
    # (we'll still use case-insensitive but keep content as-is)
    # --- 1. Layer 2 Security ---
    if not re.search(r"\bip dhcp snooping\b", content, re.IGNORECASE):
        findings.append(("DHCP Snooping Disabled", filename, "DHCP attacks possible", "Enable DHCP Snooping", "Layer 2"))

    if not re.search(r"\bip arp inspection\b", content, re.IGNORECASE):
        findings.append(("Dynamic ARP Inspection Missing", filename, "ARP spoofing possible", "Enable Dynamic ARP Inspection", "Layer 2"))

    if not re.search(r"\bswitchport port-security\b", content, re.IGNORECASE):
        findings.append(("Port Security Not Configured", filename, "MAC flooding risk", "Enable Port Security", "Layer 2"))

    # Heuristic: detect interface blocks that may not include shutdown
    try:
        # find interface blocks; heuristic: interface <name> ... (if no 'shutdown' in block, flag)
        iface_blocks = re.findall(r'(?ms)^(interface\s+\S+.*?)(?=^interface\s+\S+|\Z)', content, re.IGNORECASE)
        for block in iface_blocks:
            if not re.search(r'(?m)^\s*shutdown\b', block):
                # don't spam for each interface; append once per file as heuristic
                findings.append(("Unused Interfaces Active (heuristic)", filename, "Potential unused interfaces not administratively shutdown", "Review & administratively shutdown unused interfaces", "Layer 2"))
                break
    except Exception:
        pass

    if re.search(r"\bswitchport trunk native vlan\s+1\b", content, re.IGNORECASE):
        findings.append(("Default Native VLAN in Use", filename, "VLAN hopping risk", "Change native VLAN from 1", "Layer 2"))

    # --- 2. Access Control ---
    # Telnet detection across vty or transport input
    if re.search(r'(?mi)^\s*transport input .*telnet', content) or re.search(r'(?ms)^line vty.*?transport input .*telnet', content):
        findings.append(("Telnet Enabled", filename, "Credentials exposed in cleartext", "Disable Telnet and use SSH only", "Access Control"))

    if re.search(r"\bsnmp-server community\s+(public|private)\b", content, re.IGNORECASE):
        findings.append(("Default SNMP Community", filename, "Unauthorized SNMP access risk", "Use SNMPv3 with strong credentials", "Access Control"))

    if not re.search(r"\b(access-list|ip access-list|ip prefix-list|ipv6 access-list)\b", content, re.IGNORECASE):
        findings.append(("No ACLs Found", filename, "Unrestricted traffic flows", "Implement ACLs where needed", "Access Control"))

    # --- 3. Authentication & Authorization ---
    if not re.search(r"\baaa new-model\b", content, re.IGNORECASE):
        findings.append(("No AAA Configured", filename, "No centralized authentication", "Enable AAA (TACACS+/RADIUS)", "AAA"))

    if re.search(r'(?mi)^\s*username\s+\S+\s+(?:password|privilege)\b', content):
        findings.append(("Local User Accounts with Passwords", filename, "Local credential management; possible weak auth", "Use AAA and avoid plaintext local passwords", "AAA"))

    # --- 4. Logging & Monitoring ---
    if not re.search(r"\blogging\s+\S+", content, re.IGNORECASE):
        findings.append(("No Syslog Configured", filename, "No centralized log collection", "Configure Syslog servers", "Logging"))

    if not re.search(r"\b(ntp server|clock set|ntp peer)\b", content, re.IGNORECASE):
        findings.append(("No NTP Configured", filename, "Logs not time-synced", "Configure NTP servers", "Logging"))

    if not re.search(r"snmp-server group .* v3", content, re.IGNORECASE):
        findings.append(("SNMPv3 Not Configured", filename, "Monitoring unencrypted", "Use SNMPv3 with authentication & privacy", "Logging"))

    # --- 5. Cryptographic & Protocol Risks ---
    if re.search(r'(?mi)^\s*(service ftp|ftp server|ip ftp)\b', content):
        findings.append(("FTP Enabled", filename, "Credentials exposed in cleartext", "Disable FTP; use SFTP/SCP/FTPS", "Crypto"))

    if re.search(r'(?mi)^\s*ip http\b', content):
        findings.append(("HTTP Server Enabled", filename, "Management traffic unencrypted", "Disable HTTP; enable HTTPS (ip http secure-server)", "Crypto"))

    if not re.search(r"\bip ssh\b", content, re.IGNORECASE):
        findings.append(("SSH Not Configured", filename, "Secure remote management not enforced", "Enable SSH v2 and restrict vty to SSH", "Crypto"))

    # --- 6. Resilience & Availability ---
    if not re.search(r"\b(standby\b|vrrp\b|hsrp\b)", content, re.IGNORECASE):
        findings.append(("No First-Hop Redundancy (HSRP/VRRP)", filename, "Single point of failure for gateway", "Implement HSRP/VRRP where required", "Resilience"))

    if not re.search(r"\bstorm-control\b", content, re.IGNORECASE):
        findings.append(("No Storm Control", filename, "Broadcast/multicast flood risk", "Enable storm-control on access ports", "Resilience"))

    if not re.search(r"\bspanning-tree\b", content, re.IGNORECASE):
        findings.append(("Spanning Tree Not Configured", filename, "Switching loops possible", "Enable STP and configure root guard/portfast", "Resilience"))

    # --- 7. Configuration Management ---
    if re.search(r"\bpassword 7\b", content, re.IGNORECASE):
        findings.append(("Weak Password Encryption (Type 7)", filename, "Easily reversible encryption", "Avoid type 7; use enable secret / stronger hashes", "Config Mgmt"))

    if not re.search(r"\barchive\b", content, re.IGNORECASE):
        findings.append(("No Config Archiving", filename, "No config backup/versioning", "Enable config archive/backup/versioning", "Config Mgmt"))

    if not re.search(r"\bservice password-encryption\b", content, re.IGNORECASE):
        findings.append(("Passwords Not Encrypted", filename, "Plaintext passwords in config", "Enable 'service password-encryption' and use secrets", "Config Mgmt"))

    return findings

# ---------------------------
# Risk scoring
# ---------------------------
def get_risk_score(num_findings):
    if num_findings == 0:
        return "No Risk"
    elif num_findings <= 2:
        return "Low"
    elif num_findings <= 5:
        return "Medium"
    else:
        return "High"

# ---------------------------
# Heatmap generator
# ---------------------------
def generate_heatmap_figure(df_findings):
    """Return matplotlib figure of heatmap (devices x categories counts)."""
    if df_findings.empty:
        fig = plt.figure(figsize=(6, 3))
        plt.text(0.5, 0.5, "No data", ha='center', va='center')
        plt.axis('off')
        return fig

    pivot = pd.pivot_table(df_findings, values='Finding', index='File', columns='Category', aggfunc='count', fill_value=0)
    # ensure consistent category order
    categories_order = ["Layer 2", "Access Control", "AAA", "Logging", "Crypto", "Resilience", "Config Mgmt"]
    cols = [c for c in categories_order if c in pivot.columns] + [c for c in pivot.columns if c not in categories_order]
    pivot = pivot[cols]
    fig, ax = plt.subplots(figsize=(10, max(2, 0.35 * len(pivot.index))))
    sns.heatmap(pivot, cmap="RdYlGn_r", annot=True, fmt="d", linewidths=0.5, ax=ax)
    ax.set_title("Risk Heatmap per Category (device = row)")
    plt.tight_layout()
    return fig

# ---------------------------
# PDF generator with WIDE Device Risk Summary table
# ---------------------------
def generate_pdf_report(summary_df, df_findings, risk_counts, category_counts):
    buffer = io.BytesIO()
    
    # Use landscape orientation for wider tables
    doc = SimpleDocTemplate(buffer, pagesize=landscape(letter), 
                          leftMargin=0.3*inch, rightMargin=0.3*inch, 
                          topMargin=0.4*inch, bottomMargin=0.4*inch)
    
    styles = getSampleStyleSheet()
    
    # Create custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Title'],
        fontSize=16,
        spaceAfter=12,
        alignment=1
    )
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=12,
        spaceAfter=8
    )
    
    device_style = ParagraphStyle(
        'DeviceStyle',
        parent=styles['Heading3'],
        fontSize=10,
        spaceAfter=6,
        textColor=colors.darkblue
    )
    
    table_style = ParagraphStyle(
        'TableStyle',
        parent=styles['Normal'],
        fontSize=7,
        leading=8,
        spaceAfter=0,
        spaceBefore=0
    )
    
    elements = []

    # Title
    elements.append(Paragraph("Network Configuration Audit Report", title_style))
    elements.append(Paragraph(f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M:%SZ')}", table_style))
    elements.append(Spacer(1, 15))

    # 1. Device Risk Summary - WIDE TABLE
    elements.append(Paragraph("Device Risk Summary", heading_style))
    
    summary_data = [["Device", "Findings Count", "Risk Score"]]
    for _, row in summary_df.iterrows():
        # Use full device names - no truncation
        summary_data.append([str(row['Device']), str(row['Findings Count']), row['Risk Score']])
    
    # Calculate table width for landscape (11 inches wide - margins)
    table_width = landscape(letter)[0] - 0.6*inch  # 11 - 0.6 = 10.4 inches
    
    # Make Device column much wider to fit full names
    col_widths = [
        table_width * 0.70,  # Device - 70% of width for full names
        table_width * 0.15,  # Findings Count - 15%
        table_width * 0.15   # Risk Score - 15%
    ]
    
    summary_table = Table(summary_data, colWidths=col_widths, repeatRows=1)
    summary_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#4CAF50')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (0, -1), 'LEFT'),    # Device left-aligned
        ('ALIGN', (1, 0), (-1, -1), 'CENTER'), # Counts and Risk centered
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 9),
        ('FONTSIZE', (0, 1), (-1, -1), 8),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('WORDWRAP', (0, 0), (-1, -1), True),  # Enable word wrap for long device names
    ]))
    
    elements.append(summary_table)
    elements.append(Spacer(1, 20))

    # 2. Risk Distribution Chart
    elements.append(Paragraph("Risk Distribution (Devices by Risk Level)", heading_style))
    
    order = ["No Risk", "Low", "Medium", "High"]
    counts = [risk_counts.get(x, 0) for x in order]
    
    fig, ax = plt.subplots(figsize=(8, 4))
    bars = ax.bar(order, counts, color=["lightgrey", "lightgreen", "gold", "crimson"])
    
    for bar, count in zip(bars, counts):
        if count > 0:
            height = bar.get_height()
            ax.text(bar.get_x() + bar.get_width()/2., height + 0.1,
                   f'{count}', ha='center', va='bottom', fontsize=12, fontweight='bold')
    
    ax.set_ylabel("Number of Devices", fontsize=12)
    ax.set_title("Device Risk Distribution", fontsize=14)
    
    with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp_file:
        chart_path1 = tmp_file.name
    fig.savefig(chart_path1, bbox_inches='tight', dpi=120)
    plt.close(fig)
    
    elements.append(Image(chart_path1, width=7*inch, height=3.5*inch))
    elements.append(Spacer(1, 20))

    # 3. Findings by Category Chart
    elements.append(Paragraph("Findings by Category", heading_style))
    
    cat_names = list(category_counts.keys())
    cat_vals = list(category_counts.values())
    
    fig, ax = plt.subplots(figsize=(10, 5))
    bars = ax.bar(cat_names, cat_vals, color="steelblue")
    
    for bar, count in zip(bars, cat_vals):
        if count > 0:
            height = bar.get_height()
            ax.text(bar.get_x() + bar.get_width()/2., height + 0.1,
                   f'{count}', ha='center', va='bottom', fontsize=12, fontweight='bold')
    
    ax.set_ylabel("Number of Findings", fontsize=12)
    ax.set_title("Findings Distribution per Category", fontsize=14)
    plt.xticks(rotation=45)
    
    with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp_file:
        chart_path2 = tmp_file.name
    fig.savefig(chart_path2, bbox_inches='tight', dpi=120)
    plt.close(fig)
    
    elements.append(Image(chart_path2, width=8*inch, height=4*inch))
    elements.append(PageBreak())

    # 4. Detailed Findings with proper text wrapping
    elements.append(Paragraph("Detailed Findings", heading_style))
    
    if not df_findings.empty:
        devices = df_findings['File'].unique()
        
        for i, device in enumerate(devices):
            device_findings = df_findings[df_findings['File'] == device]
            
            elements.append(Paragraph(f"Device: {device}", device_style))
            elements.append(Spacer(1, 8))
            
            table_data = []
            
            header_style = ParagraphStyle(
                'HeaderStyle',
                parent=table_style,
                fontSize=7,
                fontName='Helvetica-Bold',
                textColor=colors.white,
                alignment=1
            )
            
            header_cells = [
                Paragraph("Category", header_style),
                Paragraph("Finding", header_style),
                Paragraph("Risk Description", header_style),
                Paragraph("Recommendation", header_style)
            ]
            table_data.append(header_cells)
            
            for _, finding in device_findings.iterrows():
                category_cell = Paragraph(str(finding['Category']), table_style)
                finding_cell = Paragraph(str(finding['Finding']), table_style)
                risk_cell = Paragraph(str(finding['RiskDesc']), table_style)
                recommendation_cell = Paragraph(str(finding['Recommendation']), table_style)
                
                table_data.append([category_cell, finding_cell, risk_cell, recommendation_cell])
            
            # Use full landscape width for detailed findings table
            table_width = landscape(letter)[0] - 0.6*inch
            col_widths = [
                table_width * 0.15,  # Category
                table_width * 0.20,  # Finding
                table_width * 0.30,  # Risk Description
                table_width * 0.35   # Recommendation
            ]
            
            device_table = Table(table_data, colWidths=col_widths, repeatRows=1)
            
            device_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2196F3')),
                ('BACKGROUND', (0, 1), (-1, -1), colors.lightgrey),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('LEFTPADDING', (0, 0), (-1, -1), 3),
                ('RIGHTPADDING', (0, 0), (-1, -1), 3),
                ('TOPPADDING', (0, 0), (-1, -1), 2),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
            ]))
            
            elements.append(device_table)
            elements.append(Spacer(1, 15))
            
            if (i + 1) % 2 == 0 and (i + 1) < len(devices):
                elements.append(PageBreak())
    
    else:
        elements.append(Paragraph("No findings to report.", table_style))

    # Build PDF
    doc.build(elements)
    pdf_bytes = buffer.getvalue()
    buffer.close()
    
    # Cleanup
    try:
        os.unlink(chart_path1)
        os.unlink(chart_path2)
    except:
        pass
    
    return pdf_bytes

# ---------------------------
# Word Document generator with wide tables
# ---------------------------
def generate_word_report(summary_df, df_findings, risk_counts, category_counts):
    doc = Document()
    
    # Set wider margins for Word document
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.3)
        section.right_margin = Inches(0.3)
    
    # Title
    title = doc.add_heading('Network Configuration Audit Report', 0)
    doc.add_paragraph(f'Generated: {datetime.utcnow().strftime("%Y-%m-%d %H:%M:%SZ")}')
    doc.add_paragraph()
    
    # 1. Device Risk Summary - WIDE TABLE
    doc.add_heading('Device Risk Summary', level=1)
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.autofit = False  # Disable autofit to control column widths
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Device'
    hdr_cells[1].text = 'Findings Count'
    hdr_cells[2].text = 'Risk Score'
    
    # Set column widths for Word table (wider first column)
    table.columns[0].width = Inches(6.0)  # Wide column for device names
    table.columns[1].width = Inches(1.5)  # Narrower for counts
    table.columns[2].width = Inches(1.5)  # Narrower for risk scores
    
    for _, row in summary_df.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(row['Device'])  # Full device name
        row_cells[1].text = str(row['Findings Count'])
        row_cells[2].text = str(row['Risk Score'])
    
    doc.add_paragraph()
    
    # 2. Risk Distribution
    doc.add_heading('Risk Distribution', level=1)
    for risk_level in ["No Risk", "Low", "Medium", "High"]:
        count = risk_counts.get(risk_level, 0)
        doc.add_paragraph(f'{risk_level}: {count} devices', style='List Bullet')
    
    doc.add_paragraph()
    
    # 3. Findings by Category
    doc.add_heading('Findings by Category', level=1)
    for category, count in category_counts.items():
        doc.add_paragraph(f'{category}: {count} findings', style='List Bullet')
    
    doc.add_paragraph()
    
    # 4. Detailed Findings
    doc.add_heading('Detailed Findings', level=1)
    
    if not df_findings.empty:
        devices = df_findings['File'].unique()
        
        for device in devices:
            doc.add_heading(f'Device: {device}', level=2)
            device_findings = df_findings[df_findings['File'] == device]
            
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            table.autofit = False
            
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Category'
            hdr_cells[1].text = 'Finding'
            hdr_cells[2].text = 'Risk Description'
            hdr_cells[3].text = 'Recommendation'
            
            # Set wider columns for Word
            table.columns[0].width = Inches(1.2)
            table.columns[1].width = Inches(2.0)
            table.columns[2].width = Inches(3.0)
            table.columns[3].width = Inches(3.0)
            
            for _, finding in device_findings.iterrows():
                row_cells = table.add_row().cells
                row_cells[0].text = str(finding['Category'])
                row_cells[1].text = str(finding['Finding'])
                row_cells[2].text = str(finding['RiskDesc'])
                row_cells[3].text = str(finding['Recommendation'])
            
            doc.add_paragraph()
    else:
        doc.add_paragraph('No findings to report.')
    
    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    word_bytes = buffer.getvalue()
    buffer.close()
    
    return word_bytes

# ---------------------------
# Main Application
# ---------------------------
def main():
    st.title("🔐 Network Config Auditor")
    
    # Create tabs with Audit Planner first
    tab1, tab2 = st.tabs(["📅 Audit Planner", "📊 Config Audit"])
    
    with tab1:
        audit_planner()
    
    with tab2:
        # Your existing config audit code
        st.markdown("""
        ### Network Configuration Audit Input
        Upload individual configuration files in text format for security assessment and compliance auditing.

        **Current Input Requirements:**
        - Text files with .txt extension
        - Supports Individual and multiple file uploads

        **Audit Outputs:**
        - Detailed vulnerability findings with remediation guidance
        - Device-level risk scoring and categorization
        - Interactive security posture dashboard
        - Risk distribution heatmaps across security categories
        - Comprehensive export formats for reporting (CSV/PDF/DOCX)
        """)

        uploaded_files = st.file_uploader(
            "Select configuration files (.txt format only) — multiple selection enabled", 
            accept_multiple_files=True, 
            type=["txt"]
        )

        if uploaded_files:
            results = []  # list of tuples: (Finding, File, RiskDesc, Recommendation, Category)
            device_summary = defaultdict(list)

            def process_file_bytes(fname, raw_bytes):
                try:
                    content = raw_bytes.decode("utf-8", errors="ignore")
                except Exception:
                    content = raw_bytes.decode("latin-1", errors="ignore")
                file_findings = audit_config(fname, content)
                for f in file_findings:
                    # f is (Finding, filename, RiskDesc, Recommendation, Category)
                    results.append(f)
                    device_summary[f[1]].append(f)
                return

            for uploaded in uploaded_files:
                name = uploaded.name
                lower = name.lower()
                # ZIP
                if lower.endswith(".zip"):
                    try:
                        with zipfile.ZipFile(io.BytesIO(uploaded.read())) as zf:
                            for inner in zf.namelist():
                                if inner.endswith("/"):
                                    continue
                                with zf.open(inner) as f:
                                    raw = f.read()
                                    process_file_bytes(inner, raw)
                    except Exception as e:
                        st.warning(f"Failed to process ZIP {name}: {e}")

                # RAR
                elif lower.endswith(".rar"):
                    try:
                        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".rar")
                        tmp.write(uploaded.read())
                        tmp.close()
                        with rarfile.RarFile(tmp.name) as rf:
                            for inner in rf.namelist():
                                if inner.endswith("/"):
                                    continue
                                with rf.open(inner) as f:
                                    raw = f.read()
                                    process_file_bytes(inner, raw)
                        try:
                            os.remove(tmp.name)
                        except Exception:
                            pass
                    except Exception as e:
                        st.warning(f"Failed to process RAR {name}: {e}")

                # Plain file (including no-extension)
                else:
                    try:
                        raw = uploaded.read()
                        process_file_bytes(name, raw)
                    except Exception as e:
                        st.warning(f"Failed to read file {name}: {e}")

            # show outputs
            if results:
                # build dataframe
                df = pd.DataFrame(results, columns=["Finding","File","RiskDesc","Recommendation","Category"])

                # Detailed findings view
                st.subheader("📋 Detailed Findings")
                st.dataframe(df[["File","Category","Finding","RiskDesc","Recommendation"]], width='stretch', height=320)

                # Device summary with risk score
                summary_rows = []
                for device, items in device_summary.items():
                    score = get_risk_score(len(items))
                    summary_rows.append((device, len(items), score))
                summary_df = pd.DataFrame(summary_rows, columns=["Device","Findings Count","Risk Score"])
                st.subheader("📊 Device Risk Summary (color-coded)")

                def color_row(r):
                    score = r["Risk Score"]
                    if score == "High":
                        return ['background-color:crimson;color:white']*3
                    if score == "Medium":
                        return ['background-color:gold;color:black']*3
                    if score == "Low":
                        return ['background-color:lightgreen;color:black']*3
                    return ['background-color:lightgrey;color:black']*3

                st.dataframe(summary_df.style.apply(lambda row: color_row(row), axis=1), width='stretch', height=220)

                # Risk distribution chart
                st.subheader("📈 Risk Distribution")
                rc = summary_df["Risk Score"].value_counts().to_dict()
                order = ["No Risk","Low","Medium","High"]
                rc_plot = [rc.get(k,0) for k in order]
                fig, ax = plt.subplots()
                bars = ax.bar(order, rc_plot, color=["lightgrey","lightgreen","gold","crimson"])
                
                for bar, count in zip(bars, rc_plot):
                    height = bar.get_height()
                    ax.text(bar.get_x() + bar.get_width()/2., height + 0.1,
                            f'{count}', ha='center', va='bottom', fontweight='bold')
                
                ax.set_ylabel("Number of Devices")
                ax.set_title("Device Risk Distribution")
                st.pyplot(fig)

                # Findings by category chart for Streamlit
                st.subheader("📊 Findings by Category")
                category_counts = df['Category'].value_counts().to_dict()
                cat_names = list(category_counts.keys())
                cat_vals = list(category_counts.values())
                fig2, ax2 = plt.subplots()
                bars2 = ax2.bar(cat_names, cat_vals, color="steelblue")
                
                for bar, count in zip(bars2, cat_vals):
                    height = bar.get_height()
                    ax2.text(bar.get_x() + bar.get_width()/2., height + 0.1,
                            f'{count}', ha='center', va='bottom', fontweight='bold')
                
                ax2.set_ylabel("Number of Findings")
                ax2.set_title("Findings Distribution per Category")
                plt.xticks(rotation=45, ha="right")
                st.pyplot(fig2)

                # Heatmap
                st.subheader("🔥 Risk Heatmap per Category")
                heatmap_fig = generate_heatmap_figure(df)
                st.pyplot(heatmap_fig)

                # Downloads: CSVs
                csv_bytes = df.to_csv(index=False).encode("utf-8")
                st.download_button("📥 Download Detailed Findings (CSV)", csv_bytes, file_name="network_detailed_findings.csv", mime="text/csv")

                csv_summary = summary_df.to_csv(index=False).encode("utf-8")
                st.download_button("📥 Download Device Summary (CSV)", csv_summary, file_name="network_device_summary.csv", mime="text/csv")

                # Management Report Generation (PDF or Word)
                st.subheader("📄 Management Report")
                
                col1, col2 = st.columns(2)
                
                with col1:
                    if st.button("Generate PDF Report"):
                        with st.spinner("Building PDF Report..."):
                            category_counts = df['Category'].value_counts().to_dict()
                            risk_counts = summary_df["Risk Score"].value_counts().to_dict()
                            pdf_bytes = generate_pdf_report(summary_df, df, risk_counts, category_counts)
                            st.success("PDF report generated successfully!")
                            st.download_button("📥 Download PDF Report", 
                                             data=pdf_bytes, 
                                             file_name="network_audit_report.pdf", 
                                             mime="application/pdf")
                
                with col2:
                    if st.button("Generate Word Report"):
                        with st.spinner("Building Word Report..."):
                            category_counts = df['Category'].value_counts().to_dict()
                            risk_counts = summary_df["Risk Score"].value_counts().to_dict()
                            word_bytes = generate_word_report(summary_df, df, risk_counts, category_counts)
                            st.success("Word report generated successfully!")
                            st.download_button("📥 Download Word Report", 
                                             data=word_bytes, 
                                             file_name="network_audit_report.docx", 
                                             mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

            else:
                st.success("✅ No findings identified in uploaded files.")
        else:
            st.info("Upload individual text files (.txt format) for configuration analysis.")

if __name__ == "__main__":
    main()